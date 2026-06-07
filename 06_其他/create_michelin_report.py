from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import urllib.parse

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0563C1')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def make_map_url(city):
    q = urllib.parse.quote(f"{city} restaurant")
    return f"https://www.google.com/maps/search/?api=1&query={q}"

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def bold_cell(cell, text, size=9):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

restaurants = {
    "ho-chi-minh": [
        {"name_en": "Ănăn Saigon", "name_cn": "安安西贡", "specialty": "Modern Vietnamese Fusion", "stars": "⭐ 1 Star", "tags": "#米其林 #fine-dining #fusion", "price": "$$$", "map": make_map_url("An An Saigon Ho Chi Minh")},
        {"name_en": "The Lunch Lady", "name_cn": "街头淑女", "specialty": "Bun bo Hue & Daily Menu", "stars": "👨‍🍳 Bourdain", "tags": "#Bourdain #legendary #street-food", "price": "$", "map": make_map_url("The Lunch Lady Ho Chi Minh")},
        {"name_en": "Nén Light", "name_cn": "轻灯餐厅", "specialty": "Contemporary Tasting Menu", "stars": "🌟 Michelin Endorsed", "tags": "#米其林 #tasting #immersive", "price": "$$$", "map": make_map_url("Nen Light Ho Chi Minh")},
        {"name_en": "Pho Le", "name_cn": "波莱河粉", "specialty": "Pho Thap Cam (Mixed Beef)", "stars": "🔥 Best in Vietnam", "tags": "#本地推荐 #pho #authentic", "price": "$", "map": make_map_url("Pho Le Ho Chi Minh")},
        {"name_en": "Banh Mi Huynh Hoa", "name_cn": "辉好法棍", "specialty": "Loaded Banh Mi", "stars": "⭐ Local Favorite", "tags": "#本地推荐 #banh-mi #loaded", "price": "$", "map": make_map_url("Banh Mi Huynh Hoa Ho Chi Minh")},
        {"name_en": "Com Nieu Sai Gon", "name_cn": "西贡砂锅饭", "specialty": "Flying Rice & Clay Pot", "stars": "👨‍🍳 Bourdain", "tags": "#Bourdain #theatrical #clay-pot", "price": "$$", "map": make_map_url("Com Nieu Sai Gon Ho Chi Minh")},
        {"name_en": "Oc Loan", "name_cn": "螺loan海鲜", "specialty": "Seafood & Shellfish", "stars": "📍 Nat Geo", "tags": "#海鲜 #shellfish #local", "price": "$$", "map": make_map_url("Oc Loan Ho Chi Minh seafood")},
        {"name_en": "Secret Garden", "name_cn": "秘密花园", "specialty": "Home Cooking", "stars": "🌿 Hidden Gem", "tags": "#隐藏餐厅 #rooftop #home-cooking", "price": "$$", "map": make_map_url("Secret Garden Ho Chi Minh restaurant")},
    ],
    "nha-trang": [
        {"name_en": "Sailing Club", "name_cn": "帆船俱乐部", "specialty": "Seafood & International", "stars": "🌊 Beachfront", "tags": "#海滨 #seafood #international", "price": "$$", "map": make_map_url("Sailing Club Nha Trang")},
        {"name_en": "Lac Canh Restaurant", "name_cn": "绿芽餐厅", "specialty": "BBQ Beef & Local", "stars": "🔥 Local BBQ", "tags": "#烧烤 #BBQ #beef", "price": "$$", "map": make_map_url("Lac Canh Restaurant Nha Trang")},
        {"name_en": "Nha Trang Seafood", "name_cn": "芽庄海鲜", "specialty": "Fresh Seafood", "stars": "🦐 Fresh Catch", "tags": "#海鲜 #fresh #local", "price": "$$", "map": make_map_url("Nha Trang Seafood Restaurant")},
        {"name_en": "Gianh Beach Restaurant", "name_cn": "迎日海滩餐厅", "specialty": "Vietnamese Seafood", "stars": "🌅 Beachside", "tags": "#海滨 #seafood #sunset", "price": "$$", "map": make_map_url("Gianh Beach Restaurant Nha Trang")},
    ],
    "dalat": [
        {"name_en": "Goc Ha Thanh", "name_cn": "江湖海鲜", "specialty": "Vietnamese & Chinese Fusion", "stars": "⭐ Local Favorite", "tags": "#本地推荐 #fusion #popular", "price": "$$", "map": make_map_url("Goc Ha Thanh Dalat restaurant")},
        {"name_en": "Da Quan", "name_cn": "大馆", "specialty": "Traditional Dalat Cuisine", "stars": "🌿 Traditional", "tags": "#传统 #local-cuisine #cozy", "price": "$$", "map": make_map_url("Da Quan Dalat restaurant")},
        {"name_en": "Lavender Dalat", "name_cn": "薰衣草餐厅", "specialty": "Vietnamese & European", "stars": "🌸 Cafe", "tags": "#咖啡厅 #european #cafe", "price": "$", "map": make_map_url("Lavender Dalat restaurant cafe")},
        {"name_en": "Dalat Night Market", "name_cn": "大叻夜市", "specialty": "Street Food & Night Market", "stars": "🌙 Night Market", "tags": "#夜市 #street-food #cheap", "price": "$", "map": make_map_url("Dalat Night Market food")},
    ],
    "taiwan": [
        {"name_en": "Din Tai Fung", "name_cn": "鼎泰丰", "specialty": "Xiaolongbao & Shanghai", "stars": "⭐⭐ Michelin 1 Star", "tags": "#米其林星 #xiaolongbao #famous", "price": "$$", "map": make_map_url("Din Tai Fung Taipei")},
        {"name_en": "Le Palais", "name_cn": "颐宫", "specialty": "Imperial Chinese Cuisine", "stars": "⭐⭐⭐ Michelin 3 Stars", "tags": "#米其林三星 #chinese #fine-dining", "price": "$$$", "map": make_map_url("Le Palais Taipei Michelin")},
        {"name_en": "Mume", "name_cn": "慕舍酒店餐厅", "specialty": "Contemporary Taiwanese", "stars": "⭐⭐ Michelin 1 Star", "tags": "#米其林星 #taiwanese #fine-dining", "price": "$$$", "map": make_map_url("Mume restaurant Taipei")},
        {"name_en": "Shin Yeh", "name_cn": "欣叶台菜", "specialty": "Traditional Taiwanese", "stars": "🌟 Local Legend", "tags": "#台菜 #traditional #local-legend", "price": "$$", "map": make_map_url("Shin Yeh restaurant Taipei")},
        {"name_en": "Jaian", "name_cn": "沾美西餐", "specialty": "Japanese Western Food", "stars": "🌟 Historic", "tags": "#西餐 #historic #buffet", "price": "$$", "map": make_map_url("Jaian restaurant Taipei")},
    ]
}

city_info = {
    "ho-chi-minh": {"name_en": "Ho Chi Minh City", "name_cn": "胡志明市", "hotel": "Sherwood Residence"},
    "nha-trang": {"name_en": "Nha Trang", "name_cn": "芽庄", "hotel": "Sunrise Nha Trang Beach Hotel"},
    "dalat": {"name_en": "DaLat", "name_cn": "大叻", "hotel": "Swiss Belresort Tuyen Lam"},
    "taiwan": {"name_en": "Taiwan", "name_cn": "台湾", "hotel": "Grand Forward Hotel"},
}

doc = Document()

# Page setup - landscape
sections = doc.sections
for section in sections:
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)

# Title
title = doc.add_heading('🍜 Michelin & Famous Restaurants Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(220, 38, 38)
    run.font.size = Pt(22)

subtitle = doc.add_paragraph('越南 · 台湾 Michelin & Famous Restaurants Guide')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(107, 114, 128)

note = doc.add_paragraph('4 Hotels | 22 Restaurants | Bilingual EN/CN |分类标签 | Map Links')
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in note.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(156, 163, 175)

# Process each city
for city_key in ["ho-chi-minh", "nha-trang", "dalat", "taiwan"]:
    city = city_info[city_key]
    city_restaurants = restaurants[city_key]

    # City header
    header = doc.add_heading(f'📍 {city["name_en"]} ({city["name_cn"]}) — {city["hotel"]}', level=1)
    for run in header.runs:
        run.font.color.rgb = RGBColor(37, 99, 235)
        run.font.size = Pt(14)

    # Table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    headers = ['Restaurant', 'Specialty', 'Rating', 'Price', 'Tags', 'Map']
    for i, h in enumerate(headers):
        bold_cell(hdr.cells[i], h)
        set_cell_shading(hdr.cells[i], '1E40AF')
        for para in hdr.cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for r in city_restaurants:
        row = table.add_row()
        cells = row.cells
        
        # Name
        cells[0].text = f'{r["name_en"]}\n{r["name_cn"]}'
        cells[0].paragraphs[0].runs[0].font.size = Pt(8)
        cells[0].paragraphs[0].runs[0].bold = True
        if len(cells[0].paragraphs[0].runs) > 1:
            cells[0].paragraphs[0].runs[1].font.size = Pt(7)
            cells[0].paragraphs[0].runs[1].font.color.rgb = RGBColor(107, 114, 128)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Specialty
        cells[1].text = r["specialty"]
        cells[1].paragraphs[0].runs[0].font.size = Pt(7)
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Stars
        cells[2].text = r["stars"]
        cells[2].paragraphs[0].runs[0].font.size = Pt(7)
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Price
        cells[3].text = r["price"]
        cells[3].paragraphs[0].runs[0].font.size = Pt(8)
        cells[3].paragraphs[0].runs[0].bold = True
        cells[3].paragraphs[0].runs[0].font.color.rgb = RGBColor(34, 197, 94)
        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Tags
        cells[4].text = r["tags"]
        cells[4].paragraphs[0].runs[0].font.size = Pt(6)
        cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(139, 92, 246)
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cells[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Map
        p = cells[5].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_hyperlink(p, '📍 Map', r["map"])
        cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph()

# Legend
legend = doc.add_heading('🏷️ Tag Legend 标签说明', level=2)
for run in legend.runs:
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 85, 99)

tags_data = [
    ("⭐ Michelin Stars", "米其林星等餐厅"),
    ("👨‍🍳 Bourdain Pick", "Bourdain推荐餐厅"),
    ("🔥 Local Favorite", "本地推荐"),
    ("🦐 Seafood", "海鲜餐厅"),
    ("🌙 Night Market", "夜市街头美食"),
    ("🌿 Hidden Gem", "隐藏版美食"),
    ("台菜 Taiwanese", "传统台菜"),
    ("☕ Cafe", "咖啡厅餐厅"),
]

tag_table = doc.add_table(rows=len(tags_data), cols=2)
tag_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (tag, desc) in enumerate(tags_data):
    row = tag_table.rows[i]
    row.cells[0].text = tag
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[0].paragraphs[0].runs[0].bold = True
    row.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(139, 92, 246)
    row.cells[1].text = desc
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(75, 85, 99)

output_path = '/Users/foongotino/Documents/RA1_Michelin_Restaurants_Guide.docx'
doc.save(output_path)
print(f"Done: {output_path}")
